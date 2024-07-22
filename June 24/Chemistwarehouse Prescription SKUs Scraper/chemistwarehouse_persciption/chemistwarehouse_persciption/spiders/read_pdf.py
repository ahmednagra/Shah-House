import glob
import os
import fitz
import pytesseract
import io
from PIL import Image
from docx import Document


class PdfDataExtraction:
    def __init__(self):
        self.file_paths = glob.glob('*.pdf')  # Adjust the path to your PDF files
        self.output_file = 'output.docx'
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Ensure this path is correct
        self.main()

    def main(self):
        doc = Document()
        for pdf_file in self.file_paths:
            print(f"\nGetting data of {os.path.basename(pdf_file)}")
            document = fitz.open(pdf_file)

            # for page_num in range(document.page_count):
            for page_num in range(0, 3):
                print('Page No :', page_num)
                page = document.load_page(page_num)
                image_list = page.get_images(full=True)

                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = document.extract_image(xref)
                    image_bytes = base_image["image"]
                    image = Image.open(io.BytesIO(image_bytes))

                    # Perform OCR for Urdu language
                    text = pytesseract.image_to_string(image, lang='urd')

                    self.add_text_to_doc(doc, text, page_num + 1)

        doc.save(self.output_file)
        print(
            f'\n\nRequired data from {len(self.file_paths)} file(s) has been updated in {os.path.basename(self.output_file)}\n')

    def add_text_to_doc(self, doc, text, page_number):
        doc.add_heading(f'Page {page_number}', level=2)
        doc.add_paragraph(text)


def main():
    PdfDataExtraction()


if __name__ == '__main__':
    main()
